import json
import re
import time
import urllib.parse
import urllib.request
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_URL = "https://lab.hebut.edu.cn/courses/python/objective/"
API_URL = BASE_URL + "qa.php"
OUT_DIR = Path(__file__).resolve().parent
RAW_JSON = OUT_DIR / "python_objective_questions.json"
OUT_DOCX = OUT_DIR / "python_objective_questions.docx"


def clean_text(value):
    value = "" if value is None else str(value)
    value = value.replace("\r\n", "\n").replace("\r", "\n").replace("\u00a0", " ")
    # Remove invisible directional marks that appear in some questions.
    value = re.sub(r"[\u202a-\u202e\u2066-\u2069]", "", value)
    value = re.sub(r"[ \t]+\n", "\n", value)
    value = re.sub(r"\n{3,}", "\n\n", value)
    return value.strip()


def fetch_batch():
    body = urllib.parse.urlencode({"action": "query"}).encode("utf-8")
    request = urllib.request.Request(
        API_URL,
        data=body,
        method="POST",
        headers={
            "User-Agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/126 Safari/537.36"
            ),
            "Referer": BASE_URL,
            "Origin": "https://lab.hebut.edu.cn",
            "X-Requested-With": "XMLHttpRequest",
            "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
        },
    )
    with urllib.request.urlopen(request, timeout=15) as response:
        payload = response.read().decode("utf-8")
    data = json.loads(payload)
    if not data.get("success"):
        raise RuntimeError(f"API returned failure: {data}")
    return data.get("data", [])


def normalize_item(item):
    return {
        "id": clean_text(item.get("ID")),
        "type": clean_text(item.get("T")),
        "field": clean_text(item.get("F")),
        "question": clean_text(item.get("Q")),
        "answer": clean_text(item.get("A")),
    }


def collect_questions(max_rounds=900, min_rounds=180, stop_no_new=180):
    questions = {}
    no_new_streak = 0
    progress = []
    for round_no in range(1, max_rounds + 1):
        added = 0
        for raw in fetch_batch():
            item = normalize_item(raw)
            key = item["id"] or f'{item["type"]}\u241f{item["question"]}\u241f{item["answer"]}'
            if item["question"] and key not in questions:
                questions[key] = item
                added += 1
        no_new_streak = 0 if added else no_new_streak + 1
        if round_no == 1 or added or round_no % 25 == 0:
            progress.append({"round": round_no, "added": added, "total": len(questions)})
            print(f"round={round_no:03d} added={added} total={len(questions)}")
        if round_no >= min_rounds and no_new_streak >= stop_no_new:
            break
        time.sleep(0.08)
    result = sorted(questions.values(), key=lambda x: (x["type"], x["field"], int(x["id"] or 0)))
    RAW_JSON.write_text(json.dumps({"source": BASE_URL, "progress": progress, "data": result}, ensure_ascii=False, indent=2), encoding="utf-8")
    return result, progress


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120):
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_multiline_paragraph(cell, text, style_name=None, bold_prefix=None):
    first = True
    for block in text.split("\n"):
        if first and len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        if style_name:
            para.style = style_name
        run = para.add_run(block)
        if bold_prefix and block.startswith(bold_prefix):
            run.bold = True
        first = False


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for attr in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, attr, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 18, 10),
        ("Heading 2", 13, "2E74B5", 14, 7),
        ("Heading 3", 12, "1F4D78", 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def build_docx(questions, progress):
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("Python 客观题题库提取")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor.from_string("0B2545")
    title.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.add_run(f"来源：{BASE_URL}\n").bold = True
    subtitle.add_run(f"提取题目数：{len(questions)}；提取方式：反复调用随机抽题接口并按题目 ID 去重。")

    by_type = Counter(q["type"] for q in questions)
    by_field = Counter(q["field"] for q in questions if q["field"])
    doc.add_heading("题目统计", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["题型", "题数", "说明"]
    for cell, text in zip(table.rows[0].cells, headers):
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for qtype, count in by_type.most_common():
        row = table.add_row().cells
        row[0].text = qtype
        row[1].text = str(count)
        row[2].text = "按题型章节整理，题目内保留原始选项和答案。"
    set_cell_margins(table)

    field_line = "；".join(f"{field} {count}题" for field, count in by_field.most_common())
    if field_line:
        doc.add_paragraph("知识点分布：" + field_line)

    grouped = defaultdict(list)
    for q in questions:
        grouped[q["type"]].append(q)

    for qtype in sorted(grouped):
        doc.add_heading(f"{qtype}（{len(grouped[qtype])}题）", level=1)
        current_field = None
        for idx, q in enumerate(grouped[qtype], 1):
            if q["field"] != current_field:
                current_field = q["field"]
                doc.add_heading(current_field or "未分类知识点", level=2)

            question_title = doc.add_paragraph()
            question_title.style = "Heading 3"
            question_title.add_run(f"{idx}. [ID {q['id']}]").bold = True

            q_para = doc.add_paragraph()
            q_para.paragraph_format.keep_with_next = True
            q_para.add_run("题目：").bold = True
            add_multiline_text(q_para, q["question"])

            ans_para = doc.add_paragraph()
            ans_para.add_run("答案：").bold = True
            add_multiline_text(ans_para, q["answer"])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "Python 客观题题库提取"
    footer.alignment = 1

    doc.save(OUT_DOCX)


def add_multiline_text(paragraph, text):
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if i:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        paragraph.add_run(line)


def main():
    questions, progress = collect_questions()
    build_docx(questions, progress)
    by_type = Counter(q["type"] for q in questions)
    print(f"saved {OUT_DOCX}")
    print(json.dumps({"total": len(questions), "by_type": by_type}, ensure_ascii=False, default=dict))


if __name__ == "__main__":
    main()
